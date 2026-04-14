"""
Document parsing utilities for PDF, DOCX, and Excel files.
Extracts text content into a unified string for evaluation.
"""
import os
import io


def parse_pdf(file_bytes):
    """Extract text from PDF bytes using PyPDF2."""
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def parse_docx(file_bytes):
    """Extract text from DOCX bytes using python-docx."""
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def parse_excel(file_bytes):
    """Extract text from Excel bytes using openpyxl."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            cell_texts = [str(c) for c in row if c is not None]
            if cell_texts:
                lines.append(" | ".join(cell_texts))
    return "\n".join(lines)


PARSERS = {
    ".pdf":  parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_excel,
    ".xls":  parse_excel,
}

SUPPORTED_EXTENSIONS = set(PARSERS.keys())


def parse_file(filename, file_bytes):
    """Parse a file by extension and return extracted text."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in PARSERS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return PARSERS[ext](file_bytes)


def load_local_file(filepath):
    """Read a local file and parse it."""
    with open(filepath, "rb") as f:
        file_bytes = f.read()
    return parse_file(os.path.basename(filepath), file_bytes)
