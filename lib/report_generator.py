"""
Report generation utilities.
Produces styled HTML and DataFrame outputs for Dataiku notebook display.
"""
import datetime
import io
import pandas as pd
from IPython.display import HTML, display


def _style_map(styler, func, **kwargs):
    """Compatibility wrapper: uses Styler.map (pandas >=2.1) or applymap (older)."""
    if hasattr(styler, "map"):
        return styler.map(func, **kwargs)
    return styler.applymap(func, **kwargs)


# ---------------------------------------------------------------------------
# Color helpers
# ---------------------------------------------------------------------------

def _score_color(score):
    if score >= 85:
        return "#2e7d32"  # green
    if score >= 70:
        return "#1565c0"  # blue
    if score >= 55:
        return "#f9a825"  # amber
    if score >= 40:
        return "#e65100"  # orange
    return "#c62828"      # red


def _verdict_color(verdict):
    if "Full Approval" in verdict:
        return "#2e7d32"
    if "Conditional" in verdict:
        return "#1565c0"
    return "#c62828"


def _bar_html(score, width=200):
    color = _score_color(score)
    filled = int(width * score / 100)
    return (
        f'<div style="background:#e0e0e0;border-radius:4px;width:{width}px;display:inline-block">'
        f'<div style="background:{color};height:16px;border-radius:4px;width:{filled}px"></div>'
        f'</div> <strong style="color:{color}">{score}</strong>'
    )


# ---------------------------------------------------------------------------
# HTML builders (return strings — used for both display and PDF export)
# ---------------------------------------------------------------------------

def _header_html(filename):
    return f"""
    <div style="background:linear-gradient(135deg,#1a237e,#283593);color:white;padding:24px 32px;border-radius:8px;margin-bottom:16px">
        <h1 style="margin:0;font-size:24px">AI / LLM Scientific Review Framework</h1>
        <p style="margin:4px 0 0;opacity:0.85">Evaluation Report &mdash; {filename}</p>
        <p style="margin:2px 0 0;opacity:0.65;font-size:12px">Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    </div>
    """


def _verdict_html(composite):
    verdict = composite["verdict"]
    score = composite["composite_score"]
    color = _verdict_color(verdict)
    interp = composite["interpretation"]
    crit = " | CRITICAL DOMAIN FAILURE" if composite["critical_failure"] else ""
    return f"""
    <div style="border:3px solid {color};border-radius:8px;padding:20px;margin-bottom:16px;text-align:center">
        <h2 style="margin:0;color:{color}">{verdict}</h2>
        <p style="font-size:36px;margin:8px 0;font-weight:bold;color:{color}">{score} / 100</p>
        <p style="margin:0;color:#555">{interp['label']} &mdash; {interp['action']}{crit}</p>
    </div>
    """


def _domain_summary_html(domain_results):
    rows = []
    for domain, data in domain_results.items():
        rows.append({
            "Domain": domain,
            "Weight": f"{int(data['weight']*100)}%",
            "Score": data["score"],
            "Interpretation": data["interpretation"]["label"],
            "Critical": "Yes" if data["is_critical"] else "",
            "Below 40": "FAIL" if data["below_threshold"] else "Pass",
        })
    df = pd.DataFrame(rows)
    s = df.style
    s = _style_map(s, lambda v: f"color: {_score_color(v)}" if isinstance(v, (int, float)) else "", subset=["Score"])
    s = _style_map(s, lambda v: "color: #c62828; font-weight: bold" if v == "FAIL" else "", subset=["Below 40"])
    return "<h3>Domain Scores</h3>" + s.hide(axis="index").to_html()


def _domain_detail_html(domain_results):
    parts = []
    for domain, data in domain_results.items():
        color = _score_color(data["score"])
        parts.append(f"""
        <div style="border-left:4px solid {color};padding:8px 16px;margin:12px 0">
            <h4 style="margin:0">{domain} — {_bar_html(data['score'], 180)}</h4>
        </div>
        """)
        rows = []
        for c in data["criteria"]:
            rows.append({
                "Criterion": c["name"],
                "Rigor": c["rigor"],
                "Score": c["score"],
                "Evidence": c.get("evidence", "")[:120] or "—",
                "Gap": c.get("gap", "")[:120] or "None",
            })
        s = pd.DataFrame(rows).style
        s = _style_map(s, lambda v: f"color: {_score_color(v)}" if isinstance(v, (int, float)) and v <= 100 else "", subset=["Score"])
        parts.append(s.hide(axis="index").to_html())
    return "".join(parts)


def _full_report_html(filename, results):
    """Build complete HTML string for one scientific evaluation report."""
    return "".join([
        _header_html(filename),
        _verdict_html(results["composite"]),
        _domain_summary_html(results["domains"]),
        "<hr><h3>Detailed Domain Evaluation</h3>",
        _domain_detail_html(results["domains"]),
    ])


# ---------------------------------------------------------------------------
# Display functions (delegate to HTML builders)
# ---------------------------------------------------------------------------

def display_header(filename):
    display(HTML(_header_html(filename)))


def display_verdict(composite):
    display(HTML(_verdict_html(composite)))


def display_domain_summary(domain_results):
    display(HTML(_domain_summary_html(domain_results)))


def display_domain_detail(domain_results):
    display(HTML(_domain_detail_html(domain_results)))


def display_full_report(filename, results):
    """Render the complete evaluation report."""
    display(HTML(_full_report_html(filename, results)))


# ---------------------------------------------------------------------------
# Word document generation and Dataiku folder export
# ---------------------------------------------------------------------------

_DATAIKU_FOLDER = "evaluation_results"


def _cell_header(cell, text):
    """Set cell text and bold the run, creating a run if needed."""
    cell.text = str(text)
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].bold = True
    else:
        para.add_run(str(text)).bold = True


def _generate_docx(filename, results, report_type="scientific"):
    """Build a python-docx Document from evaluation results and return bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title block ──────────────────────────────────────────────────────────
    title_label = "AI Artifact Evaluation" if report_type == "artifact" else "AI / LLM Scientific Review Framework"
    h = doc.add_heading(title_label, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Evaluation Report — {filename}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].italic = True

    ts = doc.add_paragraph(f"Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ts.runs:
        ts.runs[0].font.size = Pt(9)
        ts.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ── Verdict block ─────────────────────────────────────────────────────────
    composite = results["composite"]
    verdict = str(composite["verdict"])
    score = composite["composite_score"]
    interp = composite["interpretation"]
    crit = " | CRITICAL DOMAIN FAILURE" if composite["critical_failure"] else ""

    doc.add_heading("Overall Verdict", level=2)
    vp = doc.add_paragraph()
    run = vp.add_run(f"{verdict}  —  {score} / 100")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"{interp['label']} — {interp['action']}{crit}")
    doc.add_paragraph()

    # ── Domain summary table ──────────────────────────────────────────────────
    doc.add_heading("Domain Scores", level=2)
    summary_cols = ["Domain", "Weight", "Score", "Interpretation", "Critical", "Below 40"]
    tbl = doc.add_table(rows=1, cols=len(summary_cols))
    tbl.style = "Table Grid"
    for i, col in enumerate(summary_cols):
        _cell_header(tbl.rows[0].cells[i], col)

    for domain, data in results["domains"].items():
        row = tbl.add_row().cells
        row[0].text = str(domain)
        row[1].text = f"{int(data['weight'] * 100)}%"
        row[2].text = str(data["score"])
        row[3].text = str(data["interpretation"]["label"])
        row[4].text = "Yes" if data["is_critical"] else ""
        row[5].text = "FAIL" if data["below_threshold"] else "Pass"

    doc.add_paragraph()

    # ── Detailed domain criteria ──────────────────────────────────────────────
    issues_key = "issues" if report_type == "artifact" else "gap"
    issues_label = "Issues" if report_type == "artifact" else "Gap"

    doc.add_heading("Detailed Evaluation", level=2)
    for domain, data in results["domains"].items():
        doc.add_heading(f"{domain}  (score: {data['score']})", level=3)

        crit_cols = ["Criterion", "Rigor", "Score", "Evidence", issues_label]
        ctbl = doc.add_table(rows=1, cols=len(crit_cols))
        ctbl.style = "Table Grid"
        for i, col in enumerate(crit_cols):
            _cell_header(ctbl.rows[0].cells[i], col)

        for c in data["criteria"]:
            crow = ctbl.add_row().cells
            crow[0].text = str(c.get("name", ""))
            crow[1].text = str(c.get("rigor", ""))
            crow[2].text = str(c.get("score", ""))
            crow[3].text = str(c.get("evidence", "") or "—")
            crow[4].text = str(c.get(issues_key, "") or "None")

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _save_docx_to_folder(filename, results, report_type="scientific"):
    """Generate a Word document from results and save it to the Dataiku evaluation_results folder."""
    import dataiku
    try:
        docx_bytes = _generate_docx(filename, results, report_type)
        stem = filename.rsplit(".", 1)[0] if "." in filename else filename
        docx_name = f"{stem}_eval_report.docx"

        folder = dataiku.Folder(_DATAIKU_FOLDER)
        folder.upload_data(docx_name, docx_bytes)

        display(HTML(
            f'<p style="color:#2e7d32;margin:4px 0">&#10003; Word document saved to '
            f'<strong>{_DATAIKU_FOLDER}/{docx_name}</strong></p>'
        ))
    except Exception as e:
        display(HTML(
            f'<p style="color:#c62828">&#10007; Word export failed for <strong>{filename}</strong>: {e}</p>'
        ))


# ---------------------------------------------------------------------------
# DataFrame conversion
# ---------------------------------------------------------------------------

def results_to_dataframe(results):
    """Convert a single document's results to a flat DataFrame."""
    rows = []
    for domain, data in results["domains"].items():
        for c in data["criteria"]:
            rows.append({
                "domain": domain,
                "domain_weight": data["weight"],
                "domain_score": data["score"],
                "criterion": c["name"],
                "criterion_rigor": c["rigor"],
                "criterion_score": c["score"],
                "evidence": c.get("evidence", ""),
                "gap": c.get("gap", ""),
            })
    df = pd.DataFrame(rows)
    df["composite_score"] = results["composite"]["composite_score"]
    df["verdict"] = results["composite"]["verdict"]
    return df


# ---------------------------------------------------------------------------
# Batch / orchestration functions (called directly from notebook cells)
# ---------------------------------------------------------------------------

def display_all_reports(all_results):
    """Render evaluation reports for every document in all_results."""
    for filename, data in all_results.items():
        display_full_report(filename, data["results"])
        _save_docx_to_folder(filename, data["results"], report_type="scientific")
        display(HTML("<br><hr style='border:2px solid #ccc'><br>"))


def export_results(all_results):
    """
    Combine all document results into a single exportable DataFrame.
    Displays the table and returns the DataFrame.
    """
    frames = []
    for filename, data in all_results.items():
        df = results_to_dataframe(data["results"])
        df.insert(0, "document", filename)
        frames.append(df)

    if not frames:
        print("No results to export.")
        return pd.DataFrame()

    results_df = pd.concat(frames, ignore_index=True)
    display(HTML("<h3>Results DataFrame (exportable)</h3>"))
    display(results_df)
    return results_df


def display_comparison(all_results):
    """Show a side-by-side comparison table when multiple documents were evaluated."""
    if len(all_results) <= 1:
        print("Upload multiple documents to see a comparison table.")
        return

    rows = []
    for filename, data in all_results.items():
        r = data["results"]
        row = {"Document": filename}
        for domain, ddata in r["domains"].items():
            row[domain] = ddata["score"]
        row["Composite"] = r["composite"]["composite_score"]
        row["Verdict"] = r["composite"]["verdict"]
        rows.append(row)

    display(HTML("<h3>Multi-Document Comparison</h3>"))
    display(HTML(pd.DataFrame(rows).style.hide(axis="index").to_html()))


# ===========================================================================
# AI Artifact Evaluation — display and export functions
# ===========================================================================

def _artifact_header_html(filename):
    return f"""
    <div style="background:linear-gradient(135deg,#b71c1c,#c62828);color:white;padding:24px 32px;border-radius:8px;margin-bottom:16px">
        <h1 style="margin:0;font-size:24px">AI Artifact Evaluation</h1>
        <p style="margin:4px 0 0;opacity:0.85">Artifact Review &mdash; {filename}</p>
        <p style="margin:2px 0 0;opacity:0.65;font-size:12px">Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    </div>
    """


def _artifact_domain_detail_html(domain_results):
    parts = []
    for domain, data in domain_results.items():
        color = _score_color(data["score"])
        parts.append(f"""
        <div style="border-left:4px solid {color};padding:8px 16px;margin:12px 0">
            <h4 style="margin:0">{domain} — {_bar_html(data['score'], 180)}</h4>
        </div>
        """)
        rows = []
        for c in data["criteria"]:
            rows.append({
                "Criterion": c["name"],
                "Rigor": c["rigor"],
                "Score": c["score"],
                "Evidence": c.get("evidence", "")[:120] or "—",
                "Issues": c.get("issues", "")[:120] or "None",
            })
        s = pd.DataFrame(rows).style
        s = _style_map(s, lambda v: f"color: {_score_color(v)}" if isinstance(v, (int, float)) and v <= 100 else "", subset=["Score"])
        parts.append(s.hide(axis="index").to_html())
    return "".join(parts)


def _full_artifact_report_html(filename, results):
    """Build complete HTML string for one artifact evaluation report."""
    return "".join([
        _artifact_header_html(filename),
        _verdict_html(results["composite"]),
        _domain_summary_html(results["domains"]),
        "<hr><h3>Detailed Artifact Evaluation</h3>",
        _artifact_domain_detail_html(results["domains"]),
    ])


def display_artifact_report(filename, results):
    """Render a single artifact evaluation report."""
    display(HTML(_full_artifact_report_html(filename, results)))


def artifact_results_to_dataframe(results):
    """Convert a single artifact's results to a flat DataFrame."""
    rows = []
    for domain, data in results["domains"].items():
        for c in data["criteria"]:
            rows.append({
                "domain": domain,
                "domain_weight": data["weight"],
                "domain_score": data["score"],
                "criterion": c["name"],
                "criterion_rigor": c["rigor"],
                "criterion_score": c["score"],
                "evidence": c.get("evidence", ""),
                "issues": c.get("issues", ""),
            })
    df = pd.DataFrame(rows)
    df["composite_score"] = results["composite"]["composite_score"]
    df["verdict"] = results["composite"]["verdict"]
    return df


# --- Batch orchestration for artifacts ---

def display_all_artifact_reports(all_results):
    """Render artifact evaluation reports for every artifact."""
    for filename, data in all_results.items():
        display_artifact_report(filename, data["results"])
        _save_docx_to_folder(filename, data["results"], report_type="artifact")
        display(HTML("<br><hr style='border:2px solid #ccc'><br>"))


def export_artifact_results(all_results):
    """Combine all artifact results into a single exportable DataFrame."""
    frames = []
    for filename, data in all_results.items():
        df = artifact_results_to_dataframe(data["results"])
        df.insert(0, "artifact", filename)
        frames.append(df)

    if not frames:
        print("No artifact results to export.")
        return pd.DataFrame()

    results_df = pd.concat(frames, ignore_index=True)
    display(HTML("<h3>Artifact Results DataFrame (exportable)</h3>"))
    display(results_df)
    return results_df


def display_artifact_comparison(all_results):
    """Side-by-side comparison for multiple artifacts."""
    if len(all_results) <= 1:
        print("Upload multiple artifacts to see a comparison table.")
        return

    rows = []
    for filename, data in all_results.items():
        r = data["results"]
        row = {"Artifact": filename}
        for domain, ddata in r["domains"].items():
            row[domain] = ddata["score"]
        row["Composite"] = r["composite"]["composite_score"]
        row["Verdict"] = r["composite"]["verdict"]
        rows.append(row)

    display(HTML("<h3>Multi-Artifact Comparison</h3>"))
    display(HTML(pd.DataFrame(rows).style.hide(axis="index").to_html()))
